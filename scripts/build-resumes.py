from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DIR = ROOT / "content" / "resume"
OUTPUT_DIRS = [ROOT / "public" / "resume", ROOT / "dist" / "resume"]
SOURCE_FILES = [
    SOURCE_DIR / "llm-systems-engineer.json",
    SOURCE_DIR / "solution-architect.json",
]


def load_resume(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def ensure_dirs() -> None:
    for output_dir in OUTPUT_DIRS:
        output_dir.mkdir(parents=True, exist_ok=True)


def build_pdf(resume: dict, destination: Path) -> None:
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ResumeTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=20,
        leading=24,
        textColor=colors.HexColor("#0D1218"),
        alignment=TA_CENTER,
        spaceAfter=6,
    )
    headline_style = ParagraphStyle(
        "ResumeHeadline",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=11,
        leading=13,
        textColor=colors.HexColor("#0D3EC5"),
        alignment=TA_CENTER,
        spaceAfter=8,
    )
    contact_style = ParagraphStyle(
        "ResumeContact",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=9,
        leading=11,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#334253"),
        spaceAfter=3,
    )
    section_style = ParagraphStyle(
        "ResumeSection",
        parent=styles["Heading3"],
        fontName="Helvetica-Bold",
        fontSize=10,
        leading=12,
        textColor=colors.HexColor("#0D3EC5"),
        spaceBefore=8,
        spaceAfter=4,
    )
    item_title_style = ParagraphStyle(
        "ResumeItemTitle",
        parent=styles["BodyText"],
        fontName="Helvetica-Bold",
        fontSize=10,
        leading=12,
        textColor=colors.HexColor("#0D1218"),
        spaceAfter=1,
    )
    meta_style = ParagraphStyle(
        "ResumeMeta",
        parent=styles["BodyText"],
        fontName="Helvetica-Oblique",
        fontSize=8.5,
        leading=10.5,
        textColor=colors.HexColor("#334253"),
        spaceAfter=2,
    )
    body_style = ParagraphStyle(
        "ResumeBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=8.5,
        leading=10.5,
        textColor=colors.HexColor("#0D1218"),
        spaceAfter=2,
    )

    story = [
        Paragraph(resume["name"], title_style),
        Paragraph(resume["headline"], headline_style),
    ]
    for line in resume["contact_lines"]:
        story.append(Paragraph(line, contact_style))
    story.append(Spacer(1, 0.08 * inch))

    story.extend(_pdf_bullet_section("SUMMARY", resume["summary"], section_style, body_style))
    story.extend(_pdf_skill_section(resume["skill_groups"], section_style, body_style))
    story.extend(_pdf_experience_section(resume["experience"], section_style, item_title_style, meta_style, body_style))
    story.extend(_pdf_project_section(resume["projects"], section_style, item_title_style, meta_style, body_style))
    story.extend(_pdf_simple_section("EDUCATION", resume["education"], section_style, body_style))
    story.extend(_pdf_simple_section("CERTIFICATIONS", resume["certifications"], section_style, body_style))
    story.extend(_pdf_simple_section("LANGUAGES", resume["languages"], section_style, body_style))
    if resume.get("value_i_bring"):
        story.extend(_pdf_bullet_section("VALUE I BRING", resume["value_i_bring"], section_style, body_style))

    doc = SimpleDocTemplate(
        str(destination),
        pagesize=A4,
        leftMargin=0.55 * inch,
        rightMargin=0.55 * inch,
        topMargin=0.55 * inch,
        bottomMargin=0.45 * inch,
        title=resume["output_name"],
        author="Doeon Kim",
    )
    doc.build(story)


def _pdf_simple_section(title: str, lines: list[str], section_style: ParagraphStyle, body_style: ParagraphStyle) -> list:
    story = [Paragraph(title, section_style)]
    for line in lines:
        story.append(Paragraph(line, body_style))
    return story


def _pdf_bullet_section(title: str, lines: list[str], section_style: ParagraphStyle, body_style: ParagraphStyle) -> list:
    story = [Paragraph(title, section_style)]
    items = [ListItem(Paragraph(line, body_style), leftIndent=12) for line in lines]
    story.append(
        ListFlowable(
            items,
            bulletType="bullet",
            start="circle",
            bulletFontName="Helvetica",
            bulletFontSize=7,
            leftIndent=8,
        )
    )
    return story


def _pdf_skill_section(skill_groups: list[dict], section_style: ParagraphStyle, body_style: ParagraphStyle) -> list:
    story = [Paragraph("SKILLS", section_style)]
    for group in skill_groups:
        line = f"<b>{group['label']}:</b> {', '.join(group['items'])}"
        story.append(Paragraph(line, body_style))
    return story


def _pdf_experience_section(experience: list[dict], section_style: ParagraphStyle, item_title_style: ParagraphStyle, meta_style: ParagraphStyle, body_style: ParagraphStyle) -> list:
    story = [Paragraph("EXPERIENCE", section_style)]
    for item in experience:
        story.append(Paragraph(f"{item['company']} | {item['role']}", item_title_style))
        story.append(Paragraph(f"{item['location']} | {item['date']}", meta_style))
        items = [ListItem(Paragraph(line, body_style), leftIndent=12) for line in item["bullets"]]
        story.append(ListFlowable(items, bulletType="bullet", leftIndent=8, bulletFontSize=7))
    return story


def _pdf_project_section(projects: list[dict], section_style: ParagraphStyle, item_title_style: ParagraphStyle, meta_style: ParagraphStyle, body_style: ParagraphStyle) -> list:
    story = [Paragraph("SELECTED PROJECTS", section_style)]
    for project in projects:
        story.append(Paragraph(f"{project['name']} | {project['tagline']}", item_title_style))
        if project.get("subhead"):
            story.append(Paragraph(project["subhead"], meta_style))
        items = [ListItem(Paragraph(line, body_style), leftIndent=12) for line in project["bullets"]]
        story.append(ListFlowable(items, bulletType="bullet", leftIndent=8, bulletFontSize=7))
    return story


def build_docx(resume: dict, destination: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.4)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.start_type = WD_SECTION.NEW_PAGE

    styles = doc.styles
    _ensure_style(styles, "ResumeTitle", 18, bold=True)
    _ensure_style(styles, "ResumeHeadline", 11, bold=True, color="0D3EC5")
    _ensure_style(styles, "ResumeContact", 9, color="334253")
    _ensure_style(styles, "ResumeSection", 10, bold=True, color="0D3EC5")
    _ensure_style(styles, "ResumeBody", 9)
    _ensure_style(styles, "ResumeMeta", 8, italic=True, color="334253")

    title = doc.add_paragraph(style="ResumeTitle")
    title.alignment = 1
    title.add_run(resume["name"])

    headline = doc.add_paragraph(style="ResumeHeadline")
    headline.alignment = 1
    headline.add_run(resume["headline"])

    for line in resume["contact_lines"]:
        para = doc.add_paragraph(style="ResumeContact")
        para.alignment = 1
        para.add_run(line)

    _docx_bullet_section(doc, "SUMMARY", resume["summary"])
    _docx_skill_section(doc, resume["skill_groups"])
    _docx_experience_section(doc, resume["experience"])
    _docx_project_section(doc, resume["projects"])
    _docx_line_section(doc, "EDUCATION", resume["education"])
    _docx_line_section(doc, "CERTIFICATIONS", resume["certifications"])
    _docx_line_section(doc, "LANGUAGES", resume["languages"])
    if resume.get("value_i_bring"):
        _docx_bullet_section(doc, "VALUE I BRING", resume["value_i_bring"])

    doc.save(str(destination))


def _ensure_style(styles, name: str, size: int, bold: bool = False, italic: bool = False, color: str | None = None) -> None:
    if name in styles:
        style = styles[name]
    else:
        style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    font.name = "Arial"
    font.size = Pt(size)
    font.bold = bold
    font.italic = italic
    if color:
        font.color.rgb = RGBColor.from_string(color)


def _docx_heading(doc: Document, text: str) -> None:
    para = doc.add_paragraph(style="ResumeSection")
    para.space_before = Pt(8)
    para.space_after = Pt(2)
    para.add_run(text)


def _docx_bullet_section(doc: Document, title: str, lines: list[str]) -> None:
    _docx_heading(doc, title)
    for line in lines:
        para = doc.add_paragraph(style="ResumeBody")
        para.style = doc.styles["List Bullet"]
        run = para.add_run(line)
        run.font.name = "Arial"
        run.font.size = Pt(9)


def _docx_line_section(doc: Document, title: str, lines: list[str]) -> None:
    _docx_heading(doc, title)
    for line in lines:
        para = doc.add_paragraph(style="ResumeBody")
        para.add_run(line)


def _docx_skill_section(doc: Document, groups: list[dict]) -> None:
    _docx_heading(doc, "SKILLS")
    for group in groups:
        para = doc.add_paragraph(style="ResumeBody")
        label = para.add_run(f"{group['label']}: ")
        label.bold = True
        para.add_run(", ".join(group["items"]))


def _docx_experience_section(doc: Document, experience: list[dict]) -> None:
    _docx_heading(doc, "EXPERIENCE")
    for item in experience:
        title = doc.add_paragraph(style="ResumeBody")
        title_run = title.add_run(f"{item['company']} | {item['role']}")
        title_run.bold = True
        meta = doc.add_paragraph(style="ResumeMeta")
        meta.add_run(f"{item['location']} | {item['date']}")
        for bullet in item["bullets"]:
            para = doc.add_paragraph(style="ResumeBody")
            para.style = doc.styles["List Bullet"]
            para.add_run(bullet)


def _docx_project_section(doc: Document, projects: list[dict]) -> None:
    _docx_heading(doc, "SELECTED PROJECTS")
    for project in projects:
        title = doc.add_paragraph(style="ResumeBody")
        title_run = title.add_run(f"{project['name']} | {project['tagline']}")
        title_run.bold = True
        if project.get("subhead"):
            meta = doc.add_paragraph(style="ResumeMeta")
            meta.add_run(project["subhead"])
        for bullet in project["bullets"]:
            para = doc.add_paragraph(style="ResumeBody")
            para.style = doc.styles["List Bullet"]
            para.add_run(bullet)


def main() -> None:
    ensure_dirs()
    for source_path in SOURCE_FILES:
        resume = load_resume(source_path)
        for output_dir in OUTPUT_DIRS:
            build_pdf(resume, output_dir / f"{resume['output_name']}.pdf")
            build_docx(resume, output_dir / f"{resume['output_name']}.docx")


if __name__ == "__main__":
    main()
